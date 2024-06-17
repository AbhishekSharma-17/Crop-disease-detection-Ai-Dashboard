import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import os

# Set the style for the plots
sns.set(style="whitegrid")

# Define years for the data
years = np.arange(2020, 2025)

# Generate sample data for each required graph

# 1. Investment Amount Trends
investment_amounts = [1.5, 2.0, 2.8, 3.5, 4.2]  # in billion USD
number_of_deals = [50, 65, 80, 90, 100]
average_investment_per_deal = np.array(investment_amounts) / np.array(number_of_deals) * 1e3  # in million USD

# 2. Sector Preferences
sectors = ["Technology", "Healthcare", "Finance", "Energy", "Consumer Goods"]
investment_by_sector = [35, 25, 20, 10, 10]  # in percentage

# 3. Growth Metrics
revenue_growth = [10, 15, 20, 25, 30]  # percentage increase
market_expansion = [1, 2, 3, 4, 5]  # number of new markets
valuation_increase = [5, 10, 15, 20, 25]  # percentage increase

# 4. Geographical Distribution of Investments
regions = ["North America", "Europe", "Asia", "South America", "Africa"]
investment_by_region = [40, 30, 20, 5, 5]  # in percentage

# 5. Stage of Investment
investment_stages = ["Seed", "Early-Stage", "Growth-Stage"]
investment_by_stage = [20, 50, 30]  # in percentage

# 6. Investment Size Distribution
investment_sizes = np.random.normal(loc=30, scale=10, size=100)  # in million USD

# 7. Return on Investment (ROI)
roi = [5, 10, 15, 20, 25]  # in percentage

# 8. Sector Performance
sector_performance = [7, 12, 8, 15, 5]  # in percentage

# 9. Co-Investment Partnerships
co_investment_firms = ["Firm A", "Firm B", "Firm C", "Firm D"]
co_investment_counts = [10, 15, 5, 20]

# 10. Exit Strategies
exit_types = ["IPO", "Acquisition", "Buyout"]
exit_counts = [15, 30, 10]

# 11. Impact of Investments
social_impact = [3, 4, 5, 4.5, 5]  # in some units
environmental_impact = [4, 3.5, 4, 4.5, 5]  # in some units

# 12. Investment Trends by Technology
technologies = ["AI", "Blockchain", "IoT", "Fintech", "Biotech"]
investment_by_tech = [15, 10, 5, 20, 10]  # in percentage

# 13. Employee Growth in Portfolio Companies
employee_growth = [50, 100, 150, 200, 250]  # number of employees

# Create graphs and save them to a Word document

# Create a new document
from docx import Document
from docx.shared import Inches

# Initialize document
doc = Document()

# Function to add plots to document
def add_plot_to_doc(fig, doc, title):
    doc.add_heading(title, level=2)
    script_dir = os.path.dirname(os.path.abspath(__file__))
    plot_path = os.path.join(script_dir, 'temp_plot.png')
    fig.savefig(plot_path)
    doc.add_picture(plot_path, width=Inches(6))
    plt.close(fig)

# 1. Investment Amount Trends
fig, ax = plt.subplots()
ax.plot(years, investment_amounts, marker='o')
ax.set_title('Total Investment Amount by JP Morgan (2020-2024)')
ax.set_xlabel('Year')
ax.set_ylabel('Investment Amount (Billion USD)')
add_plot_to_doc(fig, doc, 'Investment Amount Trends')

fig, ax = plt.subplots()
ax.plot(years, number_of_deals, marker='o')
ax.set_title('Number of Deals by JP Morgan (2020-2024)')
ax.set_xlabel('Year')
ax.set_ylabel('Number of Deals')
add_plot_to_doc(fig, doc, 'Number of Deals by Year')

fig, ax = plt.subplots()
ax.plot(years, average_investment_per_deal, marker='o')
ax.set_title('Average Investment Amount per Deal (2020-2024)')
ax.set_xlabel('Year')
ax.set_ylabel('Average Investment Amount (Million USD)')
add_plot_to_doc(fig, doc, 'Average Investment Amount per Deal')

# 2. Sector Preferences
fig, ax = plt.subplots()
ax.bar(sectors, investment_by_sector)
ax.set_title('Sector Preferences of JP Morgan')
ax.set_xlabel('Sector')
ax.set_ylabel('Investment Percentage')
add_plot_to_doc(fig, doc, 'Sector Preferences')

# 3. Growth Metrics
fig, ax = plt.subplots()
ax.plot(years, revenue_growth, marker='o', label='Revenue Growth')
ax.plot(years, market_expansion, marker='o', label='Market Expansion')
ax.plot(years, valuation_increase, marker='o', label='Valuation Increase')
ax.set_title('Growth Metrics of Companies Supported by JP Morgan')
ax.set_xlabel('Year')
ax.set_ylabel('Growth Metrics')
ax.legend()
add_plot_to_doc(fig, doc, 'Growth Metrics')

# 4. Geographical Distribution of Investments
fig, ax = plt.subplots()
ax.bar(regions, investment_by_region)
ax.set_title('Geographical Distribution of Investments')
ax.set_xlabel('Region')
ax.set_ylabel('Investment Percentage')
add_plot_to_doc(fig, doc, 'Geographical Distribution of Investments')

# 5. Stage of Investment
fig, ax = plt.subplots()
ax.pie(investment_by_stage, labels=investment_stages, autopct='%1.1f%%', startangle=90)
ax.set_title('Stage of Investment')
add_plot_to_doc(fig, doc, 'Stage of Investment')

# 6. Investment Size Distribution
fig, ax = plt.subplots()
ax.hist(investment_sizes, bins=10)
ax.set_title('Investment Size Distribution')
ax.set_xlabel('Investment Size (Million USD)')
ax.set_ylabel('Frequency')
add_plot_to_doc(fig, doc, 'Investment Size Distribution')

# 7. Return on Investment (ROI)
fig, ax = plt.subplots()
ax.plot(years, roi, marker='o')
ax.set_title('Return on Investment (ROI)')
ax.set_xlabel('Year')
ax.set_ylabel('ROI (%)')
add_plot_to_doc(fig, doc, 'Return on Investment (ROI)')

# 8. Sector Performance
fig, ax = plt.subplots()
ax.bar(sectors, sector_performance)
ax.set_title('Sector Performance')
ax.set_xlabel('Sector')
ax.set_ylabel('Performance (%)')
add_plot_to_doc(fig, doc, 'Sector Performance')

# 9. Co-Investment Partnerships
fig, ax = plt.subplots()
ax.bar(co_investment_firms, co_investment_counts)
ax.set_title('Co-Investment Partnerships')
ax.set_xlabel('Co-Investment Firm')
ax.set_ylabel('Number of Co-Investments')
add_plot_to_doc(fig, doc, 'Co-Investment Partnerships')

# 10. Exit Strategies
fig, ax = plt.subplots()
ax.pie(exit_counts, labels=exit_types, autopct='%1.1f%%', startangle=90)
ax.set_title('Exit Strategies')
add_plot_to_doc(fig, doc, 'Exit Strategies')

# 11. Impact of Investments
fig, ax = plt.subplots()
ax.bar(years, social_impact, label='Social Impact')
ax.bar(years, environmental_impact, bottom=social_impact, label='Environmental Impact')
ax.set_title('Impact of Investments')
ax.set_xlabel('Year')
ax.set_ylabel('Impact')
ax.legend()
add_plot_to_doc(fig, doc, 'Impact of Investments')

# 12. Investment Trends by Technology
fig, ax = plt.subplots()
ax.plot(technologies, investment_by_tech, marker='o')
ax.set_title('Investment Trends by Technology')
ax.set_xlabel('Technology')
ax.set_ylabel('Investment Percentage')
add_plot_to_doc(fig, doc, 'Investment Trends by Technology')

# 13. Employee Growth in Portfolio Companies
fig, ax = plt.subplots()
ax.plot(years, employee_growth, marker='o')
ax.set_title('Employee Growth in Portfolio Companies')
ax.set_xlabel('Year')
ax.set_ylabel('Number of Employees')
add_plot_to_doc(fig, doc, 'Employee Growth in Portfolio Companies')

# Save the document
script_dir = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(script_dir, 'Investment_Analysis.docx')
doc.save(doc_path)

doc_path
