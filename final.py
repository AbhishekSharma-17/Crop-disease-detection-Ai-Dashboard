import streamlit as st
import torch
import torch.nn as nn
import torchvision.transforms as transforms
from torchvision import models
from PIL import Image
import cv2
import os
import json
import pandas as pd
import altair as alt
from streamlit_option_menu import option_menu
from docx import Document
from PIL import Image
import base64
import time


st.set_page_config(page_title="Green Sense AI", layout="wide")
# Load the logo
logo_path = os.path.join(os.path.dirname(__file__), "logo1.png")


# Function to load the image and convert it to base64
def get_base64_image(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


# Get the base64 string of the logo
logo_base64 = get_base64_image(logo_path)

# Custom CSS for logo positioning
st.markdown(
    f"""
    <style>
    .logo-container {{
        display: flex;
        align-items: center;
        justify-content: flex-start;
        padding: 20px;
    }}
    .logo {{
        width: 150px; /* Adjust the width as needed */
        height: auto;
    }}
    .title {{
        font-size: 2em;
        font-weight: bold;
        margin-left: 20px;
    }}
    </style>
    <div class="logo-container">
        <img src="data:image/png;base64,{logo_base64}" class="logo">
        <span style="font-size: 40px; font-weight:bolder; ">Welcome To Green Sense AI🌲</span>
        
    </div>
    """,
    unsafe_allow_html=True,
)


# Load LLM for chatbot
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser


# Initialize the LLM
llm = ChatGroq(
    api_key=os.getenv("GROQ_API_KEY"), model="llama3-70b-8192", temperature=0.5
)

# Define the ChatPromptTemplate for different languages
prompts = {
    "English": ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are FarmAi, an AI assistant specialized in agriculture. Answer the user's question in English.",
            ),
            ("user", "Question: {question}"),
        ]
    ),
    "Hindi": ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are FarmAi, an AI assistant specialized in agriculture. Answer the user's question in Hindi.",
            ),
            ("user", "Question: {question}"),
        ]
    ),
    "Marathi": ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are FarmAi, an AI assistant specialized in agriculture. Answer the user's question in Marathi.",
            ),
            ("user", "Question: {question}"),
        ]
    ),
    "Bhojpuri": ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are FarmAi, an AI assistant specialized in agriculture. Answer the user's question in Bhojpuri.",
            ),
            ("user", "Question: {question}"),
        ]
    ),
}

# Output parser
output_parser = StrOutputParser()

def get_disease_details(disease_class, language):
    disease_prompts = {
            "English":"Provide detailed information about the plant disease: {disease_class}.",
            "Hindi": "पौधे की बीमारी {disease_class} के बारे में विस्तृत जानकारी प्रदान करें।",
            "Marathi": "वनस्पती रोगाबद्दल तपशीलवार माहिती द्या: {disease_class}.",
            "Bhojpuri": "पौधा रोग {disease_class} के बारे में विस्तार से जानकारी दीजिए।"
        }
        
    prompt_template = disease_prompts.get(language, disease_prompts["English"])
    prompt_formatted = prompt_template.format(disease_class=disease_class)
        
    response = llm.invoke([{"role": "system", "content": prompt_formatted}])
    response_text = response.content
        
    return json.loads(json.dumps({"output": response_text}))

# Function to get chatbot response
def get_chatbot_response(question, language):
    prompt = prompts[language]
    prompt_formatted = prompt.format_prompt(question=question)
    response = llm.invoke(prompt_formatted.messages)
    response_text = response.content
    return json.loads(json.dumps({"output": response_text}))


# Set device to use GPU if available, otherwise use CPU
device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")

# Define data transformations (resizing, normalization)
transform = transforms.Compose(
    [
        transforms.Resize((256, 256)),
        transforms.ToTensor(),
        transforms.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]),
    ]
)


# Load the trained ResNet-18 model
def load_model(model_path):
    model = models.resnet18(weights=models.ResNet18_Weights.IMAGENET1K_V1)
    num_features = model.fc.in_features
    num_classes = 7  # Update this to match the number of classes in your saved model
    model.fc = nn.Linear(num_features, num_classes)
    model.load_state_dict(torch.load(model_path, map_location=device))
    model = model.to(device)
    model.eval()
    return model


# Load the model
script_dir = os.path.dirname(os.path.abspath(__file__))
model_path = os.path.join(script_dir, "models", "plant_disease_model.pth")
model = load_model(model_path)

# Define the class names (Update these to match your model's classes)
class_names = [
    "plant",
    "bercak_daun",
    "defisiensi_kalsium",
    "hangus_daun",
    "hawar_daun",
    "mosaik_vena_kuning",
    "virus_kuning_keriting",
]


# Path to user credentials file
credentials_file = "user_credentials.json"


# Load user credentials from file
def load_credentials():
    if os.path.exists(credentials_file):
        with open(credentials_file, "r") as file:
            return json.load(file)
    else:
        return {}


# Save user credentials to file
def save_credentials(credentials):
    with open(credentials_file, "w") as file:
        json.dump(credentials, file)


# Load existing user credentials
user_db = load_credentials()


# Authentication functions
def login(username, password):
    if username in user_db and user_db[username] == password:
        return True
    return False


def signup(username, password):
    if username not in user_db:
        user_db[username] = password
        save_credentials(user_db)
        return True
    return False


# Authentication state
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "signup_success" not in st.session_state:
    st.session_state.signup_success = False


def show_login_page():
    st.markdown("<h2 style='text-align: center;'>Login</h2>", unsafe_allow_html=True)
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    if st.button("Login"):
        if login(username, password):
            st.session_state.authenticated = True
            st.session_state.username = username
            st.rerun()
        else:
            st.error("Invalid credentials")


def show_signup_page():
    st.markdown("<h2 style='text-align: center;'>Signup</h2>", unsafe_allow_html=True)
    username = st.text_input("Choose a Username")
    password = st.text_input("Choose a Password", type="password")
    if st.button("Signup"):
        if signup(username, password):
            st.session_state.signup_success = True
            st.rerun()
        else:
            st.error("Signup failed")


def show_dashboard():

    # Sidebar Navigation with Icons
    with st.sidebar:
        selected = option_menu(
            menu_title="Dashboard",
            options=[
                "Home",
                "Profile",
                "Chatbot",
                "Crop Disease Detection",
                "Monitoring",
                "Crop Prediction",
                "Analysis",
                "Log Out",
            ],
            icons=[
                "house",
                "person",
                "robot",
                "camera",
                "eye",
                "graph-up-arrow",
                "bar-chart",
                "box-arrow-right",
            ],
            menu_icon="cast",
            default_index=0,
        )

    if selected == "Home":
        # st.title("Welcome to GreenSenseAI")
        def show_homepage():

            # Display an image at the top of the homepage
            # image = load_image('./logo.png')
            # st.image(image, use_column_width=True)

            # Create columns for layout
            col1, col2, col3 = st.columns(3)

            with col1:
                st.markdown(
                            """
                            <div style="background: linear-gradient(to right, rgb(182, 244, 146), rgb(51, 139, 147)); padding: 20px; border-radius: 10px; box-shadow: 2px 2px 5px rgba(0,0,0,0.1); height:250px;">
                                <h3>About Us</h3>
                                <p>GreenSenseAI is a cutting-edge AI assistant specialized in agriculture. Our mission is to provide insightful data analysis and help farmers make informed decisions.</p>
                            </div>""",
                            unsafe_allow_html=True,
                        )

            with col2:
                st.markdown(
                    """
                    <div style="height:250px; background: linear-gradient(to top, #09203f 0%, #537895 100%); padding: 20px; border-radius: 10px; box-shadow: 2px 2px 5px rgba(0,0,0,0.1);">
                        <h3>Features</h3>
                        <ul>
                            <li>Crop data analysis</li>
                            <li>Weather predictions</li>
                            <li>Pest control advice</li>
                            <li>Personalized insights</li>
                        </ul>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

            with col3:
                st.markdown(
                    """
                    <div style="height:250px;background: linear-gradient(109.6deg, rgb(0, 0, 0) 11.2%, rgb(11, 132, 145) 91.1%);
 padding: 20px; border-radius: 10px; box-shadow: 2px 2px 5px rgba(0,0,0,0.1);">
                        <h3>Get Started</h3>
                        <p>Sign up now to start using GreenSenseAI and take your agricultural practices to the next level.</p>
                        <a href="/signup" style="text-decoration: none;">
                            <button style="background-color: #4CAF50; color: white; padding: 10px 20px; border: none; border-radius: 5px; cursor: pointer;">Sign Up</button>
                        </a>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

                # Additional information section
                st.markdown(
                    """
                <div style="margin-top: 50px; padding: 20px; background: linear-gradient(to right, rgb(173, 83, 137), rgb(60, 16, 83)); border-radius: 10px; text-align: center;">
                    <h2 style="text-align:center;">Why Choose GreenSenseAI?</h2>
                    <p style="text-align:justify;">GreenSenseAI leverages advanced AI technology to provide precise and actionable insights tailored to your agricultural needs. Whether you're managing a small farm or a large agricultural enterprise, our platform adapts to deliver the best results.</p>
                </div>
                """,
                    unsafe_allow_html=True,
                )

                # Authors section at the bottom
                st.markdown(
                    """
                    <div style="margin-top: 50px; padding: 20px; background: linear-gradient(110.1deg, rgb(34, 126, 34) 2.9%, rgb(168, 251, 60) 90.3%);
 border-radius: 10px; text-align: center;">
                        <h2>Authors</h2>
                        <p>This project was developed by:</p>
                        <ul style="text-align:left;">
                            <li>Viplove Parsai</li>   
                            <li>Abhishek Sharma</li>
                            <li>Gautam Bhawsar</li>
                            <li>Ahmed Ali</li>
                        </ul>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

        show_homepage()

    elif selected == "Profile":
        st.title("GreenSenseAI Dashboard")

        # Profile Page
        st.markdown("### 👤 User Profile")
        st.write(f"**Username:** {st.session_state.username}")
        st.write("**Change Password**")
        new_password = st.text_input("New Password", type="password")
        if st.button("Update Password"):
            user_db[st.session_state.username] = new_password
            save_credentials(user_db)
            st.success("Password updated successfully.")

    elif selected == "Chatbot":
        # Chatbot Page
        st.markdown("## 🤖 FarmAi Chatbot")
        st.markdown("<hr>", unsafe_allow_html=True)

        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []

        # Dropdown for language selection
        language = st.selectbox(
            "Select Language", options=["English", "Hindi", "Marathi", "Bhojpuri"]
        )

        # Form for user input with send button below the input bar
        with st.form(key="chat_form", clear_on_submit=True):
            user_input = st.text_input(
                "Ask FarmAi about agriculture...", key="chat_input"
            )
            send_button = st.form_submit_button("Send")

        if send_button and user_input:
            st.session_state.chat_history.append(("User", user_input))
            response = get_chatbot_response(user_input, language)
            st.session_state.chat_history.append(("FarmAi", response["output"]))
            st.session_state.output = response["output"]

        # Output section for the chatbot response
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown("### 📤 Chatbot Output")
        output_container = st.container()
        if "output" in st.session_state and st.session_state["output"]:
            output_container.text_area(
                "FarmAi's Response:", value=st.session_state["output"], height=500
            )

        # Display chat history
        st.markdown("<hr>", unsafe_allow_html=True)
        with st.expander("Chat History", expanded=True):
            for speaker, message in st.session_state.chat_history:
                if speaker == "User":
                    st.markdown(
                        f"<div style='text-align: left;'><strong>User:</strong> {message}</div>",
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f"<div style='text-align: left;'><strong>FarmAi:</strong> {message}</div>",
                        unsafe_allow_html=True,
                    )

    elif selected == "Crop Disease Detection":
        
        def predict_image(image):
            image_tensor = transform(image).unsqueeze(0).to(device)
            with torch.no_grad():
                outputs = model(image_tensor)
                _, predicted = torch.max(outputs, 1)
                prediction = class_names[predicted.item()]
            return prediction
        # Crop Disease Detection Page
        st.markdown("### 📷 Real-Time Detection using Webcam")

        # Clear previous content
        for key in st.session_state.keys():
            if key.startswith("prev_"):
                del st.session_state[key]

        # Real-time detection and crop prediction
        stframe = st.empty()
        prediction_text = st.empty()
        
        st.markdown("### 🖼️ Upload Image for Detection")
        uploaded_file = st.file_uploader("Choose an image...", type=["jpg", "jpeg", "png"])

        if uploaded_file is not None:
            image = Image.open(uploaded_file)
            st.image(image, caption="Uploaded Image", use_column_width=True)

            with st.spinner("Classifying..."):
                time.sleep(2)  # Adding delay
                prediction = predict_image(image)

            prediction = predict_image(image)
            st.write(f"**Predicted Class:** {prediction}")

            
        if st.button("Get Disease Details"):
            language = st.selectbox("Select Language", ["English", "Hindi", "Marathi", "Bhojpuri"])
            disease_details = get_disease_details(prediction, language)
            st.write(disease_details["output"])
            # Provide an option to get more information via the chatbot
            st.write("For more information, ask our AI assistant:")
            question = st.text_input("Enter your question about the disease:")
            
        if st.button("Ask AI Assistant"):
                    response = get_chatbot_response(question, language)
                    st.write(response["output"])
            

        # Initialize webcam
        cap = cv2.VideoCapture(0)

        if not cap.isOpened():
            st.error("Webcam not found.")
        else:
            while True:
                ret, frame = cap.read()
                if not ret:
                    st.error("Failed to capture image from webcam.")
                    break

                # Convert frame to PIL image
                pil_image = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))

                # Transform image
                image_tensor = transform(pil_image).unsqueeze(0).to(device)

                # Perform prediction
                with torch.no_grad():
                    outputs = model(image_tensor)
                    _, predicted = torch.max(outputs, 1)
                    prediction = class_names[predicted.item()]

                # Resize frame for display
                frame_resized = cv2.resize(frame, (400, 400))

                # Display the frame and the prediction
                stframe.image(frame_resized, channels="BGR", use_column_width=True)
                prediction_text.markdown(f"### Predicted Class: **{prediction}**")

                # Exit loop on 'q' key press
                if cv2.waitKey(1) & 0xFF == ord("q"):
                    break
            
            cap.release()
            
    elif selected == "Monitoring":
        st.title("Live Feeds from Raspberry Pi 4")

        st.markdown("## **Camera Feeds**")
        st.markdown("---")

        # Create two columns for side by side feeds
        col1, col2 = st.columns(2)

        # Display the first camera feed in the first column with a frame
        with col1:
            # st.markdown("### **Camera Feed 1**")
            st.image(
                "http://192.168.31.127:5000/video_feed1",
                caption="Live Stream 1",
                use_column_width=True,
            )
            st.caption("Feed from Camera 1")
            st.markdown("---")

        # Display the second camera feed in the second column with a frame
        with col2:
            # st.markdown("### **Camera Feed 2**")
            st.image(
                "http://192.168.31.127:5000/video_feed2",
                caption="Live Stream 2",
                use_column_width=True,
            )
            st.caption("Feed from Camera 2")
            st.markdown("---")

        # Additional styling or notes
        st.markdown("### **Notes:**")
        st.markdown("1. Ensure the Raspberry Pi is connected to the network.")
        st.markdown("2. The feeds will automatically refresh to show live footage.")

    elif selected == "Crop Prediction":
        # Crop Prediction Page
        st.markdown("### 🌾 Crop Prediction")

        # Input fields for crop prediction
        moisture = st.number_input(
            "Enter the moisture level", min_value=0.0, max_value=100.0, step=0.1
        )
        humidity = st.number_input(
            "Enter the humidity level", min_value=0.0, max_value=100.0, step=0.1
        )
        temperature = st.number_input(
            "Enter the temperature (°C)", min_value=-10.0, max_value=50.0, step=0.1
        )

        # Sidebar for NPK values
        st.sidebar.header("NPK Values")
        nitrogen = st.sidebar.number_input(
            "Nitrogen level (N)", min_value=0.0, max_value=100.0, step=0.1
        )
        phosphorus = st.sidebar.number_input(
            "Phosphorus level (P)", min_value=0.0, max_value=100.0, step=0.1
        )
        potassium = st.sidebar.number_input(
            "Potassium level (K)", min_value=0.0, max_value=100.0, step=0.1
        )

        def predict_crop(
            moisture, humidity, temperature, nitrogen, phosphorus, potassium
        ):
            if (
                (temperature >= 15 and temperature <= 27)
                and (humidity >= 13 and humidity <= 14)
                and (moisture >= 21 and moisture <= 40)
            ):
                return "Rice and Wheat"
            elif (
                (temperature >= 12 and temperature <= 25)
                and (humidity >= 13 and humidity <= 15)
                and (moisture >= 21 and moisture <= 40)
            ):
                return "Maize and Pulses"
            elif (
                (temperature >= 21 and temperature <= 27)
                and (humidity >= 50 and humidity <= 60)
                and (moisture >= 45 and moisture <= 55)
            ):
                return "SugarCane"
            elif (
                (temperature >= 20 and temperature <= 31)
                and (humidity >= 40 and humidity <= 60)
                and (moisture >= 3 and moisture <= 5)
            ):
                return "Tea"
            elif (
                (temperature >= 17 and temperature <= 32)
                and (humidity >= 90 and humidity <= 95)
                and (moisture >= 95 and moisture <= 98)
            ):
                return "Cucumber"
            elif (
                (temperature >= 10 and temperature <= 21)
                and (humidity >= 90 and humidity <= 95)
                and (moisture >= 90 and moisture <= 95)
            ):
                return "Cauliflower"
            elif (
                (temperature >= 18 and temperature <= 29)
                and (humidity >= 85 and humidity <= 90)
                and (moisture >= 90 and moisture <= 95)
            ):
                return "Tomato"
            elif (
                (temperature >= 15 and temperature <= 21)
                and (humidity >= 80 and humidity <= 95)
                and (moisture >= 75 and moisture <= 80)
            ):
                return "Potato"
            elif (
                (temperature >= 24 and temperature <= 32)
                and (humidity >= 75 and humidity <= 90)
                and (moisture >= 65 and moisture <= 75)
            ):
                return "LadyFinger"
            elif (
                (temperature >= 7 and temperature <= 25)
                and (humidity >= 88 and humidity <= 95)
                and (moisture >= 90 and moisture <= 95)
            ):
                return "Cabbage"
            elif (
                (temperature >= 14 and temperature <= 32)
                and (humidity >= 70 and humidity <= 90)
                and (moisture >= 57 and moisture <= 65)
            ):
                return "Soyabean"
            elif (
                (temperature >= 13 and temperature <= 27)
                and (humidity >= 65 and humidity <= 72)
                and (moisture >= 85 and moisture <= 92)
            ):
                return "Onion"
            elif (
                (temperature >= 10 and temperature <= 29)
                and (humidity >= 80 and humidity <= 85)
                and (moisture >= 86 and moisture <= 92)
            ):
                return "Corriander"
            elif (
                (temperature >= 25 and temperature <= 35)
                and (humidity >= 12 and humidity <= 22)
                and (moisture >= 50 and moisture <= 80)
            ):
                return "Groundnut"
            elif (
                (temperature >= 32 and temperature <= 37)
                and (humidity >= 63 and humidity <= 67)
                and (moisture >= 65 and moisture <= 71)
            ):
                return "Sweet Potato"
            else:
                return "No suitable crop found"

        if st.button("Predict Crop"):
            crop = predict_crop(
                moisture, humidity, temperature, nitrogen, phosphorus, potassium
            )
            st.markdown(f"### Suitable Crop: **{crop}**")

    elif selected == "Analysis":
        # Analysis Page
        st.markdown("### 📊 Analysis")

        # Sample data for analysis
        data = {
            "Months": [
                "Jan",
                "Feb",
                "Mar",
                "Apr",
                "May",
                "Jun",
                "Jul",
                "Aug",
                "Sep",
                "Oct",
                "Nov",
                "Dec",
            ],
            "Rainfall": [78, 55, 68, 91, 102, 130, 150, 160, 120, 90, 80, 60],
            "Temperature": [15, 16, 18, 21, 24, 27, 29, 28, 25, 20, 17, 15],
            "Humidity": [75, 70, 68, 65, 60, 55, 50, 55, 60, 65, 70, 75],
        }

        df = pd.DataFrame(data)

        # Create columns for side by side charts
        col1, col2 = st.columns(2)

        # Rainfall chart with Altair
        with col1:
            st.markdown("#### Monthly Rainfall")
            rainfall_chart = (
                alt.Chart(df)
                .mark_line(point=True)
                .encode(
                    x=alt.X("Months", title="Months"),
                    y=alt.Y("Rainfall", title="Rainfall (mm)"),
                    tooltip=["Months", "Rainfall"],
                )
                .properties(title="Monthly Rainfall", width=300, height=300)
                .interactive()
            )

            st.altair_chart(rainfall_chart, use_container_width=True)
            st.markdown(
                "The rainfall chart shows the monthly distribution of rainfall. Notice the peaks in June to August, which indicate the monsoon season."
            )

        # Temperature chart with Altair
        with col2:
            st.markdown("#### Monthly Temperature")
            temperature_chart = (
                alt.Chart(df)
                .mark_line(point=True)
                .encode(
                    x=alt.X("Months", title="Months"),
                    y=alt.Y("Temperature", title="Temperature (°C)"),
                    tooltip=["Months", "Temperature"],
                )
                .properties(title="Monthly Temperature", width=300, height=300)
                .interactive()
            )

            st.altair_chart(temperature_chart, use_container_width=True)
            st.markdown(
                "The temperature chart highlights the changes in temperature over the months. The highest temperatures are observed during the summer months of June to August."
            )

        # Humidity chart with Altair
        st.markdown("#### Monthly Humidity")
        humidity_chart = (
            alt.Chart(df)
            .mark_line(point=True)
            .encode(
                x=alt.X("Months", title="Months"),
                y=alt.Y("Humidity", title="Humidity (%)"),
                tooltip=["Months", "Humidity"],
            )
            .properties(title="Monthly Humidity", width=600, height=300)
            .interactive()
        )

        st.altair_chart(humidity_chart, use_container_width=True)
        st.markdown(
            "The humidity chart displays the monthly humidity levels. Humidity decreases during the summer months and increases again in the post-monsoon period."
        )

        # Function to get analysis details from LLM
        def get_analysis_details():
            question = (
                "Provide a detailed analysis of the crop data trends over the year."
            )
            response = get_chatbot_response(question, "English")
            return response["output"]

        # Display analysis details
        if st.button("Get Analysis Details"):
            analysis_details = get_analysis_details()
            st.markdown("### Analysis Details")
            st.markdown(analysis_details)
            # Create and save analysis report as a doc file
            doc = Document()
            doc.add_heading("Analysis Report", 0)
            doc.add_paragraph(analysis_details)
            analysis_file = "analysis_report.docx"
            doc.save(analysis_file)
            st.download_button(
                label="Download Analysis Report",
                data=open(analysis_file, "rb").read(),
                file_name="analysis_report.docx",
                mime="application/octet-stream",
            )

    elif selected == "Log Out":
        st.session_state.authenticated = False
        st.rerun()


# Main logic
if st.session_state.authenticated:
    show_dashboard()
else:
    if st.session_state.signup_success:
        st.success("Signup successful! Please log in.")
        show_login_page()
    else:
        # Display tabs for Login and Signup
        tabs = st.tabs(["Login", "Signup"])
        with tabs[0]:
            show_login_page()
        with tabs[1]:
            show_signup_page()
